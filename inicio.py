# modulo para instalar las librerias necesarias
from install import *
#fin

from flask import Flask, render_template, request, redirect, url_for, flash, session
import pymysql

import os
from flask import send_file
import docx2pdf
from docx import Document
import pythoncom

from conexion import conectar

app = Flask(__name__)
app.secret_key='1234'

# equipo 7
from sql_admin import Admin
from colorama import init, Fore, Back, Style
from datetime import date

conectar()

@app.route('/verApariciones:<string:id_curso>')
def verAparicion(id_curso):
    conexion=Admin()
    # query=f"select chp.id_registro, mac.nombre, chp.lugar, chp.inicio, chp.fin, e.nombre from curso_has_aparicion chp join cursos c on c.id_curso=chp.id_curso join trabajadores e on chp.id_encargado=e.idTrabajador join modo_aplicacion_curso mac on mac.id_modo=chp.id_metodo_aplicacion where chp.id_curso={id_curso}"
    query=f"select chp.id_registro, mac.nombre, chp.lugar, c.nombre, chp.inicio, chp.fin, e.NombreTrab, chp.edicion, e.idTrabajador from curso_has_aparicion chp join cursos c on c.id_curso=chp.id_curso join trabajadores e on chp.id_encargado=e.idTrabajador join modo_aplicacion_curso mac on mac.id_modo=chp.id_metodo_aplicacion where chp.id_curso={id_curso}"
    conexion.execute(query)
    datos=conexion.Fetch()
    print(Back.YELLOW,datos,Back.RESET)
    # query=f"SELECT column_name,data_type FROM information_schema.columns  WHERE table_schema = 'rh3' AND table_name = 'curso_has_aparicion' and column_key !='PRI' order by ordinal_position"
    # cols=('modo de aplicacion','lugar','curso','inicio','fin','encargado','edicion','id encargado')
    cols=conexion.getColsNameFor('curso_has_aparicion')
    # print(f'{Back.RED}{columnas}')
    return render_template('catalogos.html',
            comentarios = datos, puestos_cursos=None,
            titulo='Aparicion',tabla_plural='APARICIONES',male=False,
                borrado=True,
                boolean=(),
                edita=True,
                id=id_curso,
                id_tabla=True,
            columnas=cols, num_columnas=len(cols),
            # columnas=('modo de aplicacion','lugar','inicio','fin','encargado'), num_columnas=5,
            acciones=0, n_acc=0)

@app.route('/Instanciar_curso:<string:id_curso>:<string:id_registro>')
def instanciarCurso(id_curso,id_registro):
    conexion=Admin()

    datos=()
    if(id_curso!='NC'):
        conexion.execute("select * from cursos where id_curso=%s"%(id_curso))
        datos+=conexion.Fetch(),
    else:
        datos+=None,

    conexion.getSQLCols("curso_has_aparicion",False)
    datos+=conexion.Fetch(),

    conexion.execute("select nombre,id_modo from modo_aplicacion_curso")
    datos+=conexion.Fetch(),

    if(id_curso!='NC'):
        query='SELECT t.idTrabajador,NombreTrab FROM trabajadores t WHERE t.idTrabajador NOT IN (SELECT che.id_empleado FROM curso_has_empleados che WHERE che.id_curso = %s)'%(id_curso)
    else:
        query=("select NombreTrab,idTrabajador from trabajadores")
        
    conexion.execute(query)
    datos+=conexion.Fetch(),
    edita=False
    #idPuesto,
    cursos=None
    if id_curso=='NC':
        # conexion.execute(f'select nomPuesto from puesto')
        conexion.execute(f'select nombre from cursos')
        cursos=conexion.Fetch()

    if( id_registro!='NR' and id_registro!='CR'):
        conexion.execute("select id_registro,lugar,inicio,fin,id_encargado,id_metodo_aplicacion  from curso_has_aparicion where id_registro=%s"%(id_registro))
        datos+=conexion.Fetch()
        edita=True
        id_registro='NR'

    return render_template("newCurso.html",comentar=datos,edita=edita,redir=id_registro,cursos=cursos)

@app.route('/ActivarCurso:<string:id_curso>:<string:id_registro>:<string:redireccion>', methods=['POST'])
def activar_curso(id_curso,id_registro,redireccion):
    print(f'{Fore.GREEN}__form recibido')
    
    Encargado=request.form['Encargado']
    MetAplicacion=request.form['MetAplicacion']
    fecha_inicio=request.form['fecha_inicio']
    fecha_fin=request.form['fecha_fin']
    lugar=request.form['lugar']

    conexion=Admin()
    activos=False

    if(id_curso=='CP'):
        NombreCurso=request.form['curso']
        print(request.form)
        conexion.execute('select id_curso from cursos where nombre="%s"'%(NombreCurso))
        id_curso=conexion.Fetch()[0][0]
        activos=True

    print(f"{Back.RED}{id_curso}")
    conexion.execute('select id_modo from modo_aplicacion_curso where nombre="%s"'%(MetAplicacion))
    id_modo=conexion.Fetch()[0][0]

    conexion.execute('select idTrabajador from trabajadores where NombreTrab="%s"'%(Encargado))
    id_encargado=conexion.Fetch()[0][0]

    if(id_registro!="NR"):
        query="update curso_has_aparicion set id_metodo_aplicacion=%s, lugar='%s', inicio='%s',fin='%s', id_encargado=%s where id_registro=%s"%(id_modo,lugar,fecha_inicio,fecha_fin,id_encargado,id_registro)
        conexion.execute(query)

        return redirect(url_for('verAparicion',id_curso=id_curso))
    # id_metodo_aplicacion,lugar,id_curso,inicio,fin,id_encargado

    conexion.execute('select veces_aparecido from cursos where id_curso=%s'%(id_curso))
    edicion=int(conexion.Fetch()[0][0])
    edicion+=1

    tupla=(id_modo,lugar,id_curso,fecha_inicio,fecha_fin,id_encargado,str(edicion))
    print(f'{Back.RED}{tupla}{Back.RESET}')

    query=('insert into curso_has_aparicion(id_metodo_aplicacion,lugar,id_curso,inicio,fin,id_encargado,edicion) values(%s,"%s","%s","%s","%s",%s,%s)'
        %tupla)
    conexion.execute(query)
    conexion.execute("update cursos set veces_aparecido=%s where id_curso=%s "%(edicion,id_curso))

    print(f"{Back.RED}{redireccion}{Back.RESET}")
    if(redireccion=='NR'):
        return redirect(url_for('verAparicion',id_curso=id_curso))

    return redirect(url_for('mostrarCatalogo',tabla='cursos',condicion='NC'))

@app.route('/catalogosEdita:<string:tabla>:<int:id_campo>')
def editarCatalogo(tabla,id_campo):
    #como 'tabla' recibe el titulo de la tabla
    conexion=Admin()
    table_name=conexion.titleToTable(tabla)

    if(table_name==None):
        return redirect("/")

    template="catalogos_edi.html"

    id_tabla=conexion.tableToId(table_name)

    conexion.execute(f"select * from {table_name}  where {id_tabla} ={id_campo}")
    dato=conexion.Fetch()

    campos= conexion.colsToString(table_name,False)[0].split(',')

    comentar=dato[0]
    # if table_name=="cursos":
    #     template="newCurso.html"
    #     comentar=dato
    return render_template(template, comentar=comentar, campo=campos, tabla=table_name, id_campo=id_campo)
# editar

# @app.route('/catalogos:<string:tabla>:<string:condicion>:<string:admin>')
@app.route('/catalogos:<string:tabla>:<string:condicion>')
# def mostrarCatalogo(tabla,condicion,admin):
def mostrarCatalogo(tabla,condicion):
    conexion=Admin()
    agregar=True
    if(conexion.existeTabla(tabla)==None ):
        # or(tabla=='curso_has_aparicion' and condicion=='NC')
        return redirect("/")

    # if(admin!='%'):
    #     admin=True
    
    permitir_borrado=True
    showId=False
    edita=True

    titulo=conexion.tableToTitle(tabla)
    id=conexion.tableToId(tabla)

    #print(f"titulo={titulo}\nid={id}")
    campos="*"
    join=""
    ci=condicion.upper()
    if(condicion=='NC' and tabla!='curso_has_aparicion'):
        condicion=""
    elif tabla=='curso_has_aparicion':
        hoy=date.today()
        print(Back.GREEN,"hoy",hoy)

        if condicion=='Activas':
            edita=False
            condicion=f'"{hoy}" between inicio and fin'

        elif condicion=='Pasadas':
            edita=False
            condicion=f' fin < "{hoy}"'
            agregar=False

        elif condicion=='Pendientes':
            condicion=f' inicio > "{hoy}"'
        else:
            condicion=''
        #el metodo comentado servira una vez que desarrolle una funcion para hacer genocidio de datos
        # join=conexion.makeJoinFor(tabla)
        campos='id_registro, modo_aplicacion_curso.nombre,lugar,cursos.nombre,inicio,fin,trabajadores.NombreTrab,edicion,trabajadores.idTrabajador'
        join='JOIN cursos ON curso_has_aparicion.id_curso=cursos.id_curso JOIN modo_aplicacion_curso ON curso_has_aparicion.id_metodo_aplicacion=modo_aplicacion_curso.id_modo JOIN trabajadores ON curso_has_aparicion.id_encargado=trabajadores.idTrabajador'
        if(condicion!=''):
            condicion="WHERE "+condicion

    if tabla=='trabajadores':
        campos='trabajadores.idTrabajador,p.nomPuesto,trabajadores.edad,trabajadores.sexo, ec.descripcion ,NombreTrab'
        join="join estado_civil ec on ec.idEstadoCivil=trabajadores.idEstadoCivil join puesto p on p.idPuesto=trabajadores.idPuesto"

    if tabla=='curso_has_empleados':
        campos='curso_has_empleados.id_registro, trabajadores.NombreTrab, curso_has_empleados.calificacion'
        join='join trabajadores on trabajadores.idTrabajador=curso_has_empleados.id_empleado'
        hoy=date.today()
        if(ci!=''):
            condicion="WHERE id_curso ="+ci
            query='select id_curso from curso_has_aparicion where id_registro=%s and fin<"%s"'%(ci,hoy)
            conexion.execute(query)
            res=conexion.Fetch()
            print(f'{Back.CYAN}{res}')
            if(res!=()):
                agregar=False
            # return f'Alto:{res}'

    query=f"select {campos} from {tabla} {join} {condicion} order by {id} asc"

    conexion.execute(query)
    datos = conexion.Fetch()

    plural,male=conexion.tablaPlural(titulo)
    titulo_columnas=conexion.getColsNameFor(tabla)
    if tabla=='curso_has_empleados':
        plural=titulo 
    plural=plural.upper()

    tipo_datos=conexion.colsToString(tabla,False)[1]

    booleanos=conexion.searchBooleanSQL(tipo_datos)

    n_columnas=len(titulo_columnas)

    n_registros=len(datos)
    puestos_cursos=()
    if(tabla=="cursos"):
        permitir_borrado=True
        titulo_columnas+="Puestos Aplicables",
        for curso in datos:
            conexion.execute(f"select p.nomPuesto from puesto_has_cursos pc join puesto p on pc.id_puesto=p.idPuesto where id_curso={curso[0]}")
            p=conexion.Fetch()
            puestos=""
            for puesto in p:
                puestos+=puesto[0]+","
            puestos_cursos+=puestos[:-1],

    return render_template("catalogos.html",
            comentarios = datos, n_registros=n_registros, puestos_cursos=puestos_cursos,
            titulo=titulo,tabla_plural=plural,male=male,
                borrado=permitir_borrado,
                boolean=booleanos,
                edita=edita,
            columnas=titulo_columnas, num_columnas=n_columnas,
            acciones=0, n_acc=0, id_tabla=showId,condicion=ci,id=None,
            agregar=agregar
            #,administrador=admin
            )
# area

@app.route('/EditaCatalogos:<string:tabla>:<int:id_campo>',methods=['POST'])
def catalogo_edita(tabla,id_campo):
    if request.method == 'POST':
        valor=request.form['valor']
        condicion='NC'

        conexion=Admin()
        campos= conexion.colsToString(tabla,True)[0]
        campos=campos.split(",")
        id_tabla=campos[0]
        index=1
        print(tabla)
        if(tabla=='curso_has_empleados'):
            index=3
            conexion.execute('select id_curso from curso_has_empleados where id_registro=%s'%(id_campo))
            condicion=conexion.Fetch()[0][0]
        otro=campos[index]

        query=('update %s set %s="%s" where %s=%s'%(tabla,otro,valor,id_tabla,id_campo))
        # print(f"{Back.RED}{query}{Back.RESET}")
        conexion.execute(query)
    return redirect(url_for('mostrarCatalogo',tabla=tabla,condicion=condicion))
#area_fedita

@app.route('/catalogoBorrar:<string:titulo>:<string:id>:<string:id_aparicion>')
def catalogo_borrar(titulo,id,id_aparicion):

    conexion=Admin()
    table_name=conexion.titleToTable(titulo)
    id_tabla=conexion.tableToId(table_name)
    print("===========")
    print(f"{Back.RED}{id_tabla}{Back.RESET}")
    print(f"{Back.RED}{table_name}{Back.RESET}")
    conexion.execute('delete from {0} where {1} = {2}'.format(table_name,id_tabla,id))

    print(titulo)
    if(titulo=='Aparicion'):
        return redirect(url_for('verAparicion',id_curso=id_aparicion ))

    return redirect(url_for('mostrarCatalogo',tabla=table_name,condicion='NC'))
# area_borrar


@app.route('/catalogosAgregar:<string:tabla_titulo>:<string:id_campo>')
def catalogo_agregar(tabla_titulo,id_campo):
    # if tabla_titulo=="Curso":
    #     return render_template("cursos_agr.html"

    conexion=Admin()
    nombre_tabla=conexion.titleToTable(tabla_titulo)
    dato,tipo_datos=conexion.colsToString(nombre_tabla,False)

    titulo_columnas=conexion.getColsNameFor(nombre_tabla)
    # titulo_columnas=['descripcion']

    # if(nombre_tabla=="cursos"):
    #     titulo_columnas=['nombre','descripcion','duracion','objetivos de aprendizaje','obligatorio']
    booleanos=conexion.searchBooleanSQL(tipo_datos)

    id_campo=''+id_campo
    id_campo=id_campo.split('_C')
    editar=True
    if(len(id_campo)>1):
        editar=False
    id_campo=(id_campo[0])

    if id_campo!="%" and editar==True:
        id_campo=int(id_campo)
        id_table=conexion.tableToId("cursos")

        conexion.execute(f"select {dato} from cursos where {id_table}={id_campo}")
        fetch =conexion.Fetch()
    else:
        fetch=()
        array=()
        for col in range(len(dato.split(","))):
            array+="",
        fetch+=array,

    if(nombre_tabla=="cursos"):
        conexion.execute("select idPuesto,nomPuesto from puesto order by idPuesto desc")
        # res=conexion.Fetch()
        titulo_columnas+="Puestos a los que va dirigido :",
        fetch+=conexion.Fetch()

    print(f'{Back.RED}__________{id_campo}')
    if(nombre_tabla=="curso_has_empleados"):
        
        if(id_campo!=''):
            query='SELECT t.idTrabajador,NombreTrab FROM trabajadores t WHERE t.idTrabajador NOT IN (SELECT che.id_empleado FROM curso_has_empleados che WHERE che.id_curso = %s)'%(id_campo)
        else:
            query=("select NombreTrab,idTrabajador from trabajadores")

        conexion.execute(query)
        # conexion.execute("select idTrabajador,NombreTrab from trabajadores order by idTrabajador desc")
        # res=conexion.Fetch() 
        fetch+=conexion.Fetch()

    cant_datos=len(titulo_columnas)
    return render_template("catalogos_agr.html",
        columna=dato,titulo=tabla_titulo
        ,datos=titulo_columnas,boolean=booleanos,cDatos=cant_datos,none=None,
        fetch=fetch,id_campo=id_campo,editar=editar
    )
#area_agregar

@app.route('/catalogoPUSH:<string:title>:<string:id_campo>', methods=['POST'])
def catalogo_aplicar_cambio(title,id_campo):
    if request.method != 'POST':
        return

    condicion='NC'
    conexion=Admin()
    table_name=conexion.titleToTable(title)

    columnas,tipos_dato=conexion.colsToString(table_name,False)
    datos =table_name,columnas,

    Nombre_columnas=conexion.getColsNameFor(table_name)
    print(f"{Back.WHITE}{Nombre_columnas}{Back.RESET}")
    print(table_name)

    if(table_name=='cursos'):
        Nombre_columnas=Nombre_columnas[:-1]

    uniones=""

    for col in range(len(Nombre_columnas)):
        column=Nombre_columnas[col]
        datos+=(request.form[column]),

        comilla=conexion.getComillas(tipos_dato[col])
        uniones+=f"{comilla}%s{comilla}"

        if col <len(Nombre_columnas)-1:
            uniones+=","

    print(uniones)
    print(f'{datos}')
    if id_campo=="%" or title=='Trabajadores en el curso':
        if(table_name=='cursos'):
            uniones+=",0"
        if(title=='Trabajadores en el curso'):
            datos=list(datos)
            conexion.execute('select idTrabajador from trabajadores where NombreTrab="%s"'%(datos[2]))
            id_traba=conexion.Fetch()[0][0]
            datos[2]=str(id_campo)+','+str(id_traba)
            if(datos[3]==''):
                datos[3]='Null'
            # datos[2]+=id_campo
            datos=tuple(datos)       
            condicion=str(id_campo)
        conexion.execute(f'insert into %s (%s) values ({uniones})'%datos)
    else:
        sets=""
        if(table_name=='cursos'):
            columnas=",".join(columnas.split(',')[:-1])

        for col in columnas.split(","):
            sets+=" %s =replace_ ,"
        sets=sets[:-1]
        # a los '%s' les mandamos el nombre del campo que van a actualizar
        sets=sets%tuple(columnas.split(","))
        #remplazamos 'replace_' para ahora guardar los valores que vamos a enviar
        sets=sets.replace("replace_","%s")

        #a los nuevos %s de la sentencia les damos las comillas si las requieren para enviar el dato respectivo correctamente
        sets=sets%tuple(uniones.split(","))
        #uniomos el layout de sets con los datos a actualizar
        sets=sets%datos[2:]
        id_tabla=conexion.tableToId(table_name)
        query=f"update {table_name} set {sets} where {id_tabla}={id_campo}"
        conexion.execute(query)

    if table_name=="cursos":
        puestos=request.form['puestos']
        puestos=puestos.replace("puesto-","").split("|")[1:]
        conexion.execute("select id_curso from cursos where nombre='%s'"%(datos[2]))
        id_curso=conexion.Fetch()[0]

        query="delete from puesto_has_cursos where id_curso=%s"%(id_curso)
        conexion.execute(query)

        values=""
        for puesto in  puestos:
            values+=f"({id_curso[0]},{puesto}),"

        values=values[:-1]
        query=f"insert into puesto_has_cursos(id_curso,id_puesto) values {values}"
        conexion.execute(query)
        # else:
        #     values=""
        #     for puesto in puestos:
        #         values=f"id_puesto={puesto}"
        #         query=f"update puesto_has_curso set {values} where id_curso={id_curso} and id_puesto={puesto}"
        #         conexion.execute(query)
    return redirect(url_for('mostrarCatalogo',tabla=table_name,condicion=condicion))
#area_fagrega

# fin equipo 7

@app.route('/')
def home():
    return render_template("home.html",inicio=True)

@app.route('/layout2')
def layout2():
    return render_template("home2.html")

@app.route('/area')
def area():
    conn =conectar()
    cursor = conn.cursor()
    cursor.execute('select idArea, descripcion from area order by idArea')
    datos = cursor.fetchall()
    global idarea
    return render_template("area.html", comentarios = datos)

@app.route('/area2')
def area2():
    conn =conectar()
    cursor = conn.cursor()
    cursor.execute('select idArea, descripcion from area order by idArea')
    datos = cursor.fetchall()
    global idarea
    return render_template("area2.html", comentarios = datos)

@app.route('/area_editar/<string:id>')
def area_editar(id):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute('select idArea, descripcion from area where idArea = %s', (id))
    dato  = cursor.fetchall()
    return render_template("area_edi.html", comentar=dato[0])

@app.route('/area_fedita/<string:id>',methods=['POST'])
def area_fedita(id):
    if request.method == 'POST':
        desc=request.form['descripcion']
        conn = conectar()
        cursor = conn.cursor()
        cursor.execute('update area set descripcion=%s where idArea=%s', (desc,id))
        conn.commit()
    return redirect(url_for('area'))

@app.route('/area_borrar/<string:id>')
def area_borrar(id):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute('delete from area where idArea = {0}'.format(id))
    conn.commit()
    return redirect(url_for('area'))

@app.route('/area_agregar')
def area_agregar():
    return render_template("area_agr.html")

@app.route('/area_fagrega', methods=['POST'])
def area_fagrega():
    if request.method == 'POST':
        desc = request.form['descripcion']
        conn =conectar()
        cursor = conn.cursor()
        cursor.execute('insert into area (descripcion) values (%s)',(desc))
        conn.commit()
    return redirect(url_for('area'))



@app.route('/puesto')
def puesto():
    conn =conectar()
    cursor = conn.cursor()

    cursor.execute('select idPuesto, nomPuesto from puesto order by idPuesto')
    datos = cursor.fetchall()

    return render_template("puesto.html", pue = datos, dat='   ', catArea = '   ', catEdoCivil = '   ', catEscolaridad = '   ',
                           catGradoAvance = '    ', catCarrera = '    ', catIdioma = ' ', catHabilidad = ' ')


@app.route('/puesto_fdetalle/<string:idP>', methods=['GET'])
def puesto_fdetalle(idP):
    conn = conectar()
    cursor = conn.cursor()

    cursor.execute('select idPuesto, nomPuesto from puesto order by idPuesto')
    datos = cursor.fetchall()

    cursor.execute('select idPuesto,codPuesto,idArea,nomPuesto,puestoJefeSup,jornada,remunMensual,prestaciones,descripcionGeneral,'
            'funciones,edad,sexo,idEstadoCivil,idEscolaridad,idGradoAvance,idCarrera,experiencia,conocimientos,manejoEquipo,'
            'reqFisicos,reqPsicologicos,responsabilidades,condicionesTrabajo from puesto where idPuesto = %s', (idP))
    dato = cursor.fetchall()

    cursor.execute('select a.idArea, a.descripcion from area a, puesto b where a.idArea = b.idArea and b.idPuesto = %s', (idP))
    datos1 = cursor.fetchall()

    cursor.execute('select a.idEstadoCivil, a.descripcion from estado_civil a, puesto b where a.idEstadoCivil = b.idEstadoCivil and b.idPuesto = %s', (idP))
    datos2 = cursor.fetchall()

    cursor.execute('select a.idEscolaridad, a.descripcion from escolaridad a, puesto b where a.idEscolaridad = b.idEscolaridad and b.idPuesto = %s', (idP))
    datos3 = cursor.fetchall()

    cursor.execute('select a.idGradoAvance, a.descripcion from grado_avance a, puesto b where a.idGradoAvance = b.idGradoAvance and b.idPuesto = %s', (idP))
    datos4 = cursor.fetchall()

    cursor.execute('select a.idCarrera, a.descripcion from carrera a, puesto b where a.idCarrera = b.idCarrera and b.idPuesto = %s', (idP))
    datos5 = cursor.fetchall()

    cursor.execute('select a.idPuesto, b.idIdioma, b.descripcion from puesto a, idioma b, puesto_has_idioma c '
                   'where a.idPuesto = c.idPuesto and b.idIdioma = c.idIdioma and a.idPuesto = %s', (idP))
    datos6 = cursor.fetchall()

    cursor.execute('select a.idPuesto, b.idHabilidad, b.descripcion from puesto a, habilidad b, puesto_has_habilidad c '
                   'where a.idPuesto = c.idPuesto and b.idHabilidad = c.idHabilidad and a.idPuesto = %s', (idP))
    datos7 = cursor.fetchall()
    return render_template("puesto.html", pue = datos, dat=dato[0], catArea=datos1[0], catEdoCivil=datos2[0], catEscolaridad=datos3[0],
                           catGradoAvance=datos4[0], catCarrera=datos5[0], catIdioma=datos6, catHabilidad=datos7)

@app.route('/puesto_borrar/<string:idP>')
def puesto_borrar(idP):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute('delete from puesto where idPuesto = %s',(idP))
    conn.commit()
    cursor.execute('delete from puesto_has_habilidad where idPuesto =%s ', (idP))
    conn.commit()
    cursor.execute('delete from puesto_has_idioma where idPuesto =%s ', (idP))
    conn.commit()
    return redirect(url_for('puesto'))


@app.route('/puesto_agrOp2')
def puesto_agrOp2():
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute('select idArea, descripcion from area ')
    datos1 = cursor.fetchall()

    cursor.execute('select idEstadoCivil, descripcion from estado_civil ')
    datos2 = cursor.fetchall()

    cursor.execute('select idEscolaridad, descripcion from escolaridad ')
    datos3 = cursor.fetchall()

    cursor.execute('select idGradoAvance, descripcion from grado_avance ')
    datos4 = cursor.fetchall()

    cursor.execute('select idCarrera, descripcion from carrera ')
    datos5 = cursor.fetchall()

    cursor.execute('select idIdioma, descripcion from idioma ')
    datos6 = cursor.fetchall()

    cursor.execute('select idHabilidad, descripcion from habilidad ')
    datos7 = cursor.fetchall()

    return render_template("puesto_agrOp2.html", catArea=datos1, catEdoCivil=datos2, catEscolaridad=datos3,
                           catGradoAvance=datos4, catCarrera=datos5, catIdioma=datos6, catHabilidad=datos7)


@app.route('/puesto_fagrega2', methods=['POST'])
def puesto_fagrega():
    if request.method == 'POST':

        if  'idArea' in request.form:
            idAr = request.form['idArea']
        else:
            idAr = '1'
        if 'idEstadoCivil' in request.form:
            idEC = request.form['idEstadoCivil']
        else:
            idEC = '1'
        if 'idEscolaridad' in request.form:
            idEs = request.form['idEscolaridad']
        else:
            idEs = '1'
        if 'idGradoAvance' in request.form:
            idGA = request.form['idGradoAvance']
        else:
            idGA = '1'
        if 'idCarrera' in request.form:
            idCa = request.form['idCarrera']
        else:
            idCa = '1'
        if 'sexo' in request.form:
            sex = request.form['sexo']
        else:
            sex = '1'
        codP = request.form['codPuesto']
        nomP = request.form['nomPuesto']
        pueJ = request.form['puestoJefeSup']
        jorn = request.form['jornada']
        remu = request.form['remunMensual']
        pres = request.form['prestaciones']
        desc = request.form['descripcionGeneral']
        func = request.form['funciones']
        eda = request.form['edad']
        expe = request.form['experiencia']
        cono = request.form['conocimientos']
        manE = request.form['manejoEquipo']
        reqF = request.form['reqFisicos']
        reqP = request.form['reqPsicologicos']
        resp = request.form['responsabilidades']
        conT = request.form['condicionesTrabajo']


    conn = conectar()
    cursor = conn.cursor()
    cursor.execute(
    'insert into puesto (codPuesto,idArea,nomPuesto,puestoJefeSup,jornada,remunMensual,prestaciones,descripcionGeneral,'
    'funciones,edad,sexo,idEstadoCivil,idEscolaridad,idGradoAvance,idCarrera,experiencia,conocimientos,manejoEquipo,'
    'reqFisicos,reqPsicologicos,responsabilidades,condicionesTrabajo) values (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)',
    (codP, idAr, nomP, pueJ, jorn, remu, pres, desc, func, eda, sex, idEC, idEs, idGA, idCa, expe, cono, manE, reqF,
     reqP, resp, conT))
    conn.commit()

    cursor.execute('select idPuesto from puesto where idPuesto=(select max(idPuesto) from puesto) ')
    dato = cursor.fetchall()
    idpue = dato[0]
    idP = idpue[0]

    cursor.execute('select count(*) from idioma ')
    dato = cursor.fetchall()
    nidio = dato[0]
    ni = nidio[0] + 1

    for i in range(1, ni):
        idio = 'i' + str(i)
        if idio in request.form:
            cursor.execute('insert into puesto_has_idioma(idPuesto,idIdioma) values (%s,%s)', (idP, i))
            conn.commit()

    cursor.execute('select count(*) from habilidad ')
    dato = cursor.fetchall()
    nhab = dato[0]
    nh =nhab[0]+1

    for i in range(1,nh):
        habi = 'h' + str(i)
        if habi in request.form:
            cursor.execute('insert into puesto_has_habilidad(idPuesto,idHabilidad) values (%s,%s)', (idP,i))
            conn.commit()

    return redirect(url_for('puesto'))



@app.route('/puesto_editar/<string:idP>')
def puesto_editar(idP):
    conn = conectar()
    cursor = conn.cursor()

    cursor.execute('select idPuesto,codPuesto,idArea,nomPuesto,puestoJefeSup,jornada,remunMensual,prestaciones,descripcionGeneral,'
        'funciones,edad,sexo,idEstadoCivil,idEscolaridad,idGradoAvance,idCarrera,experiencia,conocimientos,manejoEquipo,'
        'reqFisicos,reqPsicologicos,responsabilidades,condicionesTrabajo from puesto where idPuesto = %s', (idP))
    dato = cursor.fetchall()

    cursor.execute('select idArea, descripcion from area ')
    datos1 = cursor.fetchall()

    cursor.execute('select idEstadoCivil, descripcion from estado_civil ')
    datos2 = cursor.fetchall()

    cursor.execute('select idEscolaridad, descripcion from escolaridad ')
    datos3 = cursor.fetchall()

    cursor.execute('select idGradoAvance, descripcion from grado_avance ')
    datos4 = cursor.fetchall()

    cursor.execute('select idCarrera, descripcion from carrera ')
    datos5 = cursor.fetchall()

    cursor.execute('select idIdioma, descripcion from idioma ')
    datos6 = cursor.fetchall()

    cursor.execute('select idHabilidad, descripcion from habilidad ')
    datos7 = cursor.fetchall()

    cursor.execute('select a.idArea, a.descripcion from area a, puesto b where a.idArea = b.idArea and b.idPuesto = %s', (idP))
    datos11 = cursor.fetchall()

    cursor.execute('select a.idEstadoCivil, a.descripcion from estado_civil a, puesto b where a.idEstadoCivil = b.idEstadoCivil and b.idPuesto = %s',(idP))
    datos12 = cursor.fetchall()

    cursor.execute('select a.idEscolaridad, a.descripcion from escolaridad a, puesto b where a.idEscolaridad = b.idEscolaridad and b.idPuesto = %s',(idP))
    datos13 = cursor.fetchall()

    cursor.execute('select a.idGradoAvance, a.descripcion from grado_avance a, puesto b where a.idGradoAvance = b.idGradoAvance and b.idPuesto = %s',(idP))
    datos14 = cursor.fetchall()

    cursor.execute('select a.idCarrera, a.descripcion from carrera a, puesto b where a.idCarrera = b.idCarrera and b.idPuesto = %s', (idP))
    datos15 = cursor.fetchall()

    cursor.execute('select a.idPuesto, b.idIdioma, b.descripcion from puesto a, idioma b, puesto_has_idioma c '
                   'where a.idPuesto = c.idPuesto and b.idIdioma = c.idIdioma and a.idPuesto = %s', (idP))
    datos16 = cursor.fetchall()

    cursor.execute('select a.idPuesto, b.idHabilidad, b.descripcion from puesto a, habilidad b, puesto_has_habilidad c '
                   'where a.idPuesto = c.idPuesto and b.idHabilidad = c.idHabilidad and a.idPuesto = %s', (idP))
    datos17 = cursor.fetchall()


    return render_template("puesto_edi.html", dat=dato[0], catArea=datos1, catEdoCivil=datos2, catEscolaridad=datos3,
                           catGradoAvance=datos4, catCarrera=datos5, catIdioma=datos6, catHabilidad=datos7,
                           Area=datos11[0], EdoCivil=datos12[0], Escolaridad=datos13[0], GradoAvance=datos14[0],
                           Carrera=datos15[0], Idioma=datos16, Habilidad=datos17)


@app.route('/puesto_fedita/<string:idP>', methods=['POST'])
def puesto_fedita(idP):
    if request.method == 'POST':
        codP = request.form['codPuesto']
        idAr = request.form['idArea']
        nomP = request.form['nomPuesto']
        pueJ = request.form['puestoJefeSup']
        jorn = request.form['jornada']
        remu = request.form['remunMensual']
        pres = request.form['prestaciones']
        desc = request.form['descripcionGeneral']
        func = request.form['funciones']
        eda = request.form['edad']
        sex = request.form['sexo']
        idEC = request.form['idEstadoCivil']
        idEs = request.form['idEscolaridad']
        idGA = request.form['idGradoAvance']
        idCa = request.form['idCarrera']
        expe = request.form['experiencia']
        cono = request.form['conocimientos']
        manE = request.form['manejoEquipo']
        reqF = request.form['reqFisicos']
        reqP = request.form['reqPsicologicos']
        resp = request.form['responsabilidades']
        conT = request.form['condicionesTrabajo']

    conn = conectar()
    cursor = conn.cursor()

    cursor.execute('update puesto set codPuesto = %s, idArea = %s, nomPuesto = %s, puestoJefeSup = %s, jornada = %s, '
                   'remunMensual = %s, prestaciones = %s, descripcionGeneral = %s, funciones = %s, edad = %s, sexo = %s, '
                   'idEstadoCivil = %s, idEscolaridad = %s, idGradoAvance = %s, idCarrera = %s, experiencia = %s, '
                   'conocimientos = %s, manejoEquipo = %s, reqFisicos = %s, reqPsicologicos = %s, responsabilidades = %s, '
                   'condicionesTrabajo = %s where idPuesto = %s', (codP, idAr, nomP, pueJ, jorn, remu, pres, desc, func, eda,
                   sex, idEC, idEs, idGA, idCa, expe, cono, manE, reqF, reqP, resp, conT, idP))
    conn.commit()

    cursor.execute('delete from puesto_has_habilidad where idPuesto =%s ', (idP))
    conn.commit()
    cursor.execute('delete from puesto_has_idioma where idPuesto =%s ', (idP))
    conn.commit()

    cursor.execute('select count(*) from idioma ')
    dato = cursor.fetchall()
    nidio = dato[0]
    ni = nidio[0] + 1

    for i in range(1, ni):
        idio = 'i' + str(i)
        if idio in request.form:
            cursor.execute('insert into puesto_has_idioma(idPuesto,idIdioma) values (%s,%s)', (idP, i))
            conn.commit()

    cursor.execute('select count(*) from habilidad ')
    dato = cursor.fetchall()
    nhab = dato[0]
    nh = nhab[0] + 1

    for i in range(1, nh):
        habi = 'h' + str(i)
        if habi in request.form:
            cursor.execute('insert into puesto_has_habilidad(idPuesto,idHabilidad) values (%s,%s)', (idP, i))
            conn.commit()
    return redirect(url_for('puesto'))

#Inicia modificaciones de los catalogos con CRUD

#Carrera

@app.route('/carrera')
def carrera():
    conn =conectar()
    cursor = conn.cursor()
    cursor.execute('select idCarrera, descripcion from carrera order by idCarrera')
    datos = cursor.fetchall()
    return render_template("carrera.html", comentarios = datos)

@app.route('/carrera2')
def carrera2():
    conn =conectar()
    cursor = conn.cursor()
    cursor.execute('select idCarrera, descripcion from carrera order by idCarrera')
    datos = cursor.fetchall()
    return render_template("carrera2.html", comentarios = datos)

@app.route('/carrera_editar/<string:id>')
def carrera_editar(id):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute('select idCarrera, descripcion from carrera where idCarrera = %s', (id))
    dato  = cursor.fetchall()
    return render_template("carrera_edi.html", comentar=dato[0])

@app.route('/carrera_fedita/<string:id>',methods=['POST'])
def carrera_fedita(id):
    if request.method == 'POST':
        desc=request.form['descripcion']
        conn = conectar()
        cursor = conn.cursor()
        cursor.execute('update carrera set descripcion=%s where idCarrera=%s', (desc,id))
        conn.commit()
    return redirect(url_for('carrera'))

@app.route('/carrera_borrar/<string:id>')
def carrera_borrar(id):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute('delete from carrera where idCarrera = {0}'.format(id))
    conn.commit()
    return redirect(url_for('carrera'))

@app.route('/carrera_agregar')
def carrera_agregar():
    return render_template("carrera_agr.html")

@app.route('/carrera_fagrega', methods=['POST'])
def carrera_fagrega():
    if request.method == 'POST':
        desc = request.form['descripcion']
        conn =conectar()
        cursor = conn.cursor()
        cursor.execute('insert into carrera (descripcion) values (%s)',(desc))
        conn.commit()
    return redirect(url_for('carrera'))

# Escolaridad

@app.route('/escolaridad')
def escolaridad():
    conn =conectar()
    cursor = conn.cursor()
    cursor.execute('select idEscolaridad, descripcion from escolaridad order by idEscolaridad')
    datos = cursor.fetchall()
    return render_template("escolaridad.html", comentarios = datos)

@app.route('/escolaridad2')
def escolaridad2():
    conn =conectar()
    cursor = conn.cursor()
    cursor.execute('select idEscolaridad, descripcion from escolaridad order by idEscolaridad')
    datos = cursor.fetchall()
    return render_template("escolaridad2.html", comentarios = datos)

@app.route('/escolaridad_editar/<string:id>')
def escolaridad_editar(id):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute('select idEscolaridad, descripcion from escolaridad where idEscolaridad = %s', (id))
    dato  = cursor.fetchall()
    return render_template("escolaridad_edi.html", comentar=dato[0])

@app.route('/escolaridad_fedita/<string:id>',methods=['POST'])
def escolaridad_fedita(id):
    if request.method == 'POST':
        desc=request.form['descripcion']
        conn = conectar()
        cursor = conn.cursor()
        cursor.execute('update escolaridad set descripcion=%s where idEscolaridad=%s', (desc,id))
        conn.commit()
    return redirect(url_for('escolaridad'))

@app.route('/escolaridad_borrar/<string:id>')
def escolaridad_borrar(id):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute('delete from escolaridad where idEscolaridad = {0}'.format(id))
    conn.commit()
    return redirect(url_for('escolaridad'))

@app.route('/escolaridad_agregar')
def escolaridad_agregar():
    return render_template("escolaridad_agr.html")

@app.route('/escolaridad_fagrega', methods=['POST'])
def escolaridad_fagrega():
    if request.method == 'POST':
        desc = request.form['descripcion']
        conn =conectar()
        cursor = conn.cursor()
        cursor.execute('insert into escolaridad (descripcion) values (%s)',(desc))
        conn.commit()
    return redirect(url_for('escolaridad'))

# Estado Civil

@app.route('/estado_civil')
def estado_civil():
    conn =conectar()
    cursor = conn.cursor()
    cursor.execute('select idEstadoCivil, descripcion from estado_civil order by idEstadoCivil')
    datos = cursor.fetchall()
    return render_template("estado_civil.html", comentarios = datos)

@app.route('/estado_civil2')
def estado_civil2():
    conn =conectar()
    cursor = conn.cursor()
    cursor.execute('select idEstadoCivil, descripcion from estado_civil order by idEstadoCivil')
    datos = cursor.fetchall()
    return render_template("estado_civil2.html", comentarios = datos)

@app.route('/estado_civil_editar/<string:id>')
def estado_civil_editar(id):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute('select idEstadoCivil, descripcion from estado_civil where idEstadoCivil = %s', (id))
    dato  = cursor.fetchall()
    return render_template("estado_civil_edi.html", comentar=dato[0])

@app.route('/estado_civil_fedita/<string:id>',methods=['POST'])
def estado_civil_fedita(id):
    if request.method == 'POST':
        desc=request.form['descripcion']
        conn = conectar()
        cursor = conn.cursor()
        cursor.execute('update estado_civil set descripcion=%s where idEstadoCivil=%s', (desc,id))
        conn.commit()
    return redirect(url_for('estado_civil'))

@app.route('/estado_civil_borrar/<string:id>')
def estado_civil_borrar(id):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute('delete from estado_civil where idEstadoCivil = {0}'.format(id))
    conn.commit()
    return redirect(url_for('estado_civil'))

@app.route('/estado_civil_agregar')
def estado_civil_agregar():
    return render_template("estado_civil_agr.html")

@app.route('/estado_civil_fagrega', methods=['POST'])
def estado_civil_fagrega():
    if request.method == 'POST':
        desc = request.form['descripcion']
        conn =conectar()
        cursor = conn.cursor()
        cursor.execute('insert into estado_civil (descripcion) values (%s)',(desc))
        conn.commit()
    return redirect(url_for('estado_civil'))

#Grado de avance

@app.route('/grado_avance')
def grado_avance():
    conn =conectar()
    cursor = conn.cursor()
    cursor.execute('select idGradoAvance, descripcion from grado_avance order by idGradoAvance')
    datos = cursor.fetchall()
    return render_template("grado_avance.html", comentarios = datos)

@app.route('/grado_avance2')
def grado_avance2():
    conn =conectar()
    cursor = conn.cursor()
    cursor.execute('select idGradoAvance, descripcion from grado_avance order by idGradoAvance')
    datos = cursor.fetchall()
    return render_template("grado_avance2.html", comentarios = datos)

@app.route('/grado_avance_editar/<string:id>')
def grado_avance_editar(id):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute('select idGradoAvance, descripcion from grado_avance where idGrado_avance = %s', (id))
    dato  = cursor.fetchall()
    return render_template("grado_avance_edi.html", comentar=dato[0])

@app.route('/grado_avance_fedita/<string:id>',methods=['POST'])
def grado_avance_fedita(id):
    if request.method == 'POST':
        desc=request.form['descripcion']
        conn = conectar()
        cursor = conn.cursor()
        cursor.execute('update grado_avance set descripcion=%s where idGrado_avance=%s', (desc,id))
        conn.commit()
    return redirect(url_for('grado_avance'))

@app.route('/grado_avance_borrar/<string:id>')
def grado_avance_borrar(id):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute('delete from grado_avance where idGradoAvance = {0}'.format(id))
    conn.commit()
    return redirect(url_for('grado_avance'))

@app.route('/grado_avance_agregar')
def grado_avance_agregar():
    return render_template("grado_avance_agr.html")

@app.route('/grado_avance_fagrega', methods=['POST'])
def grado_avance_fagrega():
    if request.method == 'POST':
        desc = request.form['descripcion']
        conn =conectar()
        cursor = conn.cursor()
        cursor.execute('insert into grado_avance (descripcion) values (%s)',(desc))
        conn.commit()
    return redirect(url_for('grado_avance'))

#Habilidad

@app.route('/habilidad')
def habilidad():
    conn =conectar()
    cursor = conn.cursor()
    cursor.execute('select idHabilidad, descripcion from habilidad order by idHabilidad')
    datos = cursor.fetchall()
    return render_template("habilidad.html", comentarios = datos)

@app.route('/habilidad2')
def habilidad2():
    conn =conectar()
    cursor = conn.cursor()
    cursor.execute('select idHabilidad, descripcion from habilidad order by idHabilidad')
    datos = cursor.fetchall()
    return render_template("habilidad2.html", comentarios = datos)

@app.route('/habilidad_editar/<string:id>')
def habilidad_editar(id):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute('select idHabilidad, descripcion from habilidad where idHabilidad = %s', (id))
    dato  = cursor.fetchall()
    return render_template("habilidad_edi.html", comentar=dato[0])

@app.route('/habilidad_fedita/<string:id>',methods=['POST'])
def habilidad_fedita(id):
    if request.method == 'POST':
        desc=request.form['descripcion']
        conn = conectar()
        cursor = conn.cursor()
        cursor.execute('update habilidad set descripcion=%s where idHabilidad=%s', (desc,id))
        conn.commit()
    return redirect(url_for('habilidad'))

@app.route('/habilidad_borrar/<string:id>')
def habilidad_borrar(id):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute('delete from habilidad where idHabilidad = {0}'.format(id))
    conn.commit()
    return redirect(url_for('habilidad'))

@app.route('/habilidad_agregar')
def habilidad_agregar():
    return render_template("habilidad_agr.html")

@app.route('/habilidad_fagrega', methods=['POST'])
def habilidad_fagrega():
    if request.method == 'POST':
        desc = request.form['descripcion']
        conn =conectar()
        cursor = conn.cursor()
        cursor.execute('insert into habilidad (descripcion) values (%s)',(desc))
        conn.commit()
    return redirect(url_for('habilidad'))

#Idioma

@app.route('/idioma')
def idioma():
    conn =conectar()
    cursor = conn.cursor()
    cursor.execute('select idIdioma, descripcion from idioma order by idIdioma')
    datos = cursor.fetchall()
    return render_template("idioma.html", comentarios = datos)

@app.route('/idioma2')
def idioma2():
    conn =conectar()
    cursor = conn.cursor()
    cursor.execute('select idIdioma, descripcion from idioma order by idIdioma')
    datos = cursor.fetchall()
    return render_template("idioma2.html", comentarios = datos)

@app.route('/idioma_editar/<string:id>')
def idioma_editar(id):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute('select idIdioma, descripcion from idioma where idIdioma = %s', (id))
    dato  = cursor.fetchall()
    return render_template("idioma_edi.html", comentar=dato[0])

@app.route('/idioma_fedita/<string:id>',methods=['POST'])
def idioma_fedita(id):
    if request.method == 'POST':
        desc=request.form['descripcion']
        conn = conectar()
        cursor = conn.cursor()
        cursor.execute('update idioma set descripcion=%s where idIdioma=%s', (desc,id))
        conn.commit()
    return redirect(url_for('idioma'))

@app.route('/idioma_borrar/<string:id>')
def idioma_borrar(id):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute('delete from idioma where idIdioma = {0}'.format(id))
    conn.commit()
    return redirect(url_for('idioma'))

@app.route('/idioma_agregar')
def idioma_agregar():
    return render_template("idioma_agr.html")

@app.route('/idioma_fagrega', methods=['POST'])
def idioma_fagrega():
    if request.method == 'POST':
        desc = request.form['descripcion']
        conn =conectar()
        cursor = conn.cursor()
        cursor.execute('insert into idioma (descripcion) values (%s)',(desc))
        conn.commit()
    return redirect(url_for('idioma'))
#Inicio del módulo de contratación###############################################################
#Equipo 3 Ramiro#
@app.route('/login_form')
def login_form():
    return render_template('inicio_emergente.html')

@app.route('/login', methods=['POST'])
def login():
    conn =conectar()
    nombre = request.form['nombre']
    contraseña = request.form['contraseña']

    with conn.cursor() as cursor:
        sql = "SELECT * FROM admins WHERE nombre = %s AND contraseña = %s"
        cursor.execute(sql, (nombre, contraseña))
        admin = cursor.fetchone()

    if admin:
        return redirect(url_for('layout2',inicio=True))

    else:
        return "Acceso denegado. Por favor, verifica tu nombre de usuario y contraseña. <a href='/'>Volver al inicio</a>"
    


#CRUD Vacantes
@app.route('/vacante2')
def vacante2():
        conn =conectar()
        cursor = conn.cursor()
        cursor.execute('select idVacante, conseVR, fuenteCandidato, inicioFechaPublic, finFechaPublic, publicada, observaciones, candidatoSelecc, fechaContratacion, idRequisicion, idPuesto from vacante vacante where estado_vacante="autorizada" order by idVacante')
        datos=cursor.fetchall()
        return render_template("vacantes2.html", comentarios=datos)

@app.route('/vacante')
def vacante():
    conn =conectar()
    cursor = conn.cursor()
    cursor.execute('select idVacante, conseVR, fuenteCandidato, inicioFechaPublic, finFechaPublic, publicada, observaciones, candidatoSelecc, fechaContratacion, idRequisicion, idPuesto from vacante where estado_vacante="autorizada" order by idVacante')
    datos = cursor.fetchall()
    return render_template("vacantes.html", comentarios = datos)

@app.route('/vacante_editar/<string:id>')
def vacante_editar(id):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute('select idVacante, conseVR, fuenteCandidato, inicioFechaPublic, finFechaPublic, publicada, observaciones, candidatoSelecc, fechaContratacion, idRequisicion, idPuesto from vacante where idVacante = %s', (id))
    dato  = cursor.fetchall()
    return render_template("vacantes_edi.html", comentar=dato[0])

@app.route('/vacante_fedita/<string:id>',methods=['POST'])
def vacante_fedita(id):
    if request.method == 'POST':
        conse=request.form['conseVR']
        fuente=request.form['fuenteCandidato']
        fechai=request.form['inicioFechaPublic']
        fechaf=request.form['finFechaPublic']
        publi=request.form['publicada']
        obser=request.form['observaciones']
        candi=request.form['candidatoSelecc']
        fechac=request.form['fechaContratacion']
        requi=request.form['idRequisicion']
        puesto=request.form['idPuesto']

        conn = conectar()
        cursor = conn.cursor()
        cursor.execute('update vacante set conseVR=%s, fuenteCandidato=%s, inicioFechaPublic=%s, finFechaPublic=%s, publicada=%s, observaciones=%s, candidatoSelecc=%s, fechaContratacion=%s, idRequisicion=%s, idPuesto=%s where idVacante=%s', (conse, fuente, fechai, fechaf, publi, obser, candi, fechac,requi,puesto, id))
        conn.commit()
    return redirect(url_for('vacante'))

@app.route('/vacante_borrar/<string:id>')
def vacante_borrar(id):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute('delete from vacante where idVacante = {0}'.format(id))
    conn.commit()
    return redirect(url_for('vacantes'))

@app.route('/vacante_agregar')
def vacante_agregar():
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute('SELECT DISTINCT descripcion FROM publi')  
    opciones_publicada = [row[0] for row in cursor.fetchall()] 
    cursor.execute('SELECT DISTINCT nombre FROM candidato')  
    opciones_publicada2 = [row[0] for row in cursor.fetchall()]  
    cursor.execute('SELECT DISTINCT idRequisicion FROM requisicion')  
    opciones_publicada3 = [row[0] for row in cursor.fetchall()]  
    cursor.execute('SELECT DISTINCT nomPuesto FROM puesto')  
    opciones_publicada4 = [row[0] for row in cursor.fetchall()]  
    conn.close()
    return render_template("vacantes_agr.html", opciones_publicada=opciones_publicada,opciones_publicada2=opciones_publicada2,opciones_publicada3=opciones_publicada3,opciones_publicada4=opciones_publicada4)

@app.route('/vacante_fagrega', methods=['POST'])
def vacante_fagrega():
    conn =conectar()
    cursor = conn.cursor()
    publi2=0
    puesto2=0
    conse=0
    estado="solicitada" #estado_vacante
    if request.method == 'POST':
        fuente=request.form['fuenteCandidato']
        fechai=request.form['inicioFechaPublic']
        fechaf=request.form['finFechaPublic']
        publi=request.form['publicada']
        if publi=='Publicada':
            publi2==1
        else:
            publi2==2
        obser=request.form['observaciones']
        fechac=request.form['fechaContratacion']
        requi=request.form['idRequisicion']
        puesto=request.form['idPuesto']
        cursor.execute('select idPuesto from puesto where nomPuesto=%s', (puesto))
        puesto2=cursor.fetchone()

        cursor.execute('insert into vacante (conseVR,fuenteCandidato, inicioFechaPublic, finFechaPublic, publicada, observaciones, fechaContratacion, idRequisicion, idPuesto,estado_Vacante) values (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)',(conse,fuente, fechai, fechaf, publi2, obser, fechac,requi,puesto2,estado))
        conn.commit()
        cursor.execute('select idVacante from vacante where estado_vacante=%s', (estado))
        nombress=cursor.fetchall()
        return render_template('autoricvac.html', titulv=nombress)

@app.route('/conssolic', methods=['POST', 'GET'])
def conssolic():
    if request.method == 'POST':
        sol = request.form['idVA']
        conn = conectar()
        cursor = conn.cursor()
        if sol is not None:
            cursor.execute('SELECT * FROM vacante WHERE idVacante=%s', (sol,))
            solicitexist = cursor.fetchone()
            session['sol'] = sol
            print(solicitexist)
            print (sol)
        else:
            print("solicitud no encontrada")
            return render_template('vacsaut.html')
        
    return render_template('autoricvac2.html', solicitexist=solicitexist)

@app.route('/aut', methods=['POST', 'GET'])
def vacsaut():
    estadoaut="autorizada"
    sol = session.get('sol')
    print("Valor de sol en vacsaut:", sol)
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute('update vacante set estado_vacante=%s where idVacante=%s', (estadoaut, sol))
    conn.commit()
    return render_template('layout2.html', sol=sol)



@app.route('/noaut', methods=['POST', 'GET'])
def vacsnoaut():
    conn=conectar()
    cursor=conn.cursor()
    sol = session.get('sol')
    cursor.execute('update vacante set estado_vacante=%s where idVacante=%s', ('no autorizada', sol))
    conn.commit()
    return render_template('layout2.html')





#CRUD Candidatos
@app.route('/candidatos')
def candidatos():
    conn =conectar()
    cursor = conn.cursor()
    cursor.execute('select idCandidato, idVacante, idRequisicion, idPuesto, CURP, RFC, nombre, domCalle, domNumExtInt, domColonia, tel1, tel2, correoE, edad, sexo, idEstadoCivil idEscolaridad, idGradoAvance, idCarrera, entrevSelecReq, entrevSelecPresen, entrevSelecresult, evalMedicaReq, evalMedicaPresen, evalMedicaResult, evalPsicolgReq, evalPsicologPresen, evalPsicologResult, evalPsicometReq, evalPsicometPresene, evalPsicometResult, evalTecnicaReq, evalTecnicaPresen, evalPsicometResult, evalConocReq, evalConocPresen, evalConocResult, entrevFinalReq, entrevFinalPresen, entrevFinalResul from candidato order by idCandidato')
    datos = cursor.fetchall()
    return render_template("candidatos.html", comentar = datos)

@app.route('/login_form2')
def login_form2():
    return render_template('candidatour.html')

@app.route('/login2', methods=['POST'])
def login2():
    conn =conectar()
    cursor = conn.cursor()
    nombre = request.form['nombre']
    cursor.execute('select idCandidato, idVacante, idRequisicion, idPuesto, CURP, RFC, nombre, domCalle, domNumExtInt, domColonia, tel1, tel2, correoE, edad, sexo, idEstadoCivil idEscolaridad, idGradoAvance, idCarrera, entrevSelecReq, entrevSelecPresen, entrevSelecresult, evalMedicaReq, evalMedicaPresen, evalMedicaResult, evalPsicolgReq, evalPsicologPresen, evalPsicologResult, evalPsicometReq, evalPsicometPresene, evalPsicometResult, evalTecnicaReq, evalTecnicaPresen, evalPsicometResult, evalConocReq, evalConocPresen, evalConocResult, entrevFinalReq, entrevFinalPresen, entrevFinalResul from candidato  where nombre= %s', (nombre))
    datos = cursor.fetchall()
    return render_template("candidatou.html", comentar = datos)    

@app.route('/candidatos_editar/<string:id>')
def candidatos_editar(id):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute('select idCandidato, CURP, RFC, nombre, domCalle, domNumExtInt, domColonia, tel1, tel2, correoE, edad, sexo, idEstadoCivil idEscolaridad, idGradoAvance, idCarrera, entrevSelecReq, entrevSelecPresen, entrevSelecresult, evalMedicaReq, evalMedicaPresen, evalMedicaResult, evalPsicolgReq, evalPsicologPresen, evalPsicologResult, evalPsicometReq, evalPsicometPresene, evalPsicometResult, evalTecnicaReq, evalTecnicaPresen, evalPsicometResult, evalConocReq, evalConocPresen, evalConocResult, entrevFinalReq, entrevFinalPresen, entrevFinalResul from candidato where idCandidato= %s', (id))
    dato  = cursor.fetchall()
    return render_template("candidatos_edi.html", comentar=dato[0])

@app.route('/candidatos_fedita/<string:id>',methods=['POST'])
def candidatos_fedita(id):
    conn = conectar()
    cursor = conn.cursor()
    if request.method == 'POST':
        curp=request.form['CURP']
        rfc=request.form['RFC']
        nom=request.form['nombre']
        calle=request.form['domCalle']
        numex=request.form['domNumExtInt']
        col=request.form['domColonia']
        tel1=request.form['tel1']
        tel2=request.form['tel2']
        correo=request.form['correoE']
        edad=request.form['edad']
        sexo=request.form['sexo']
        idec=request.form['idEstadoCivil']
        ide=request.form['idEscolaridad']
        idga=request.form['idGradoAvance']
        idc=request.form['idCarrera']
        entsreq=request.form['entrevSelecReq']
        entsp=request.form['entrevSelecPresen']
        entsres=request.form['entrevSelecResult']
        evmedreq=request.form['evalMedicaReq']
        evmedp=request.form['evalMedicaPresen']
        evmedres=request.form['evalMedicaResul']
        evpsireq=request.form['evalPsicolgReq']
        evpsip=request.form['evalPsicologPresen']
        evpsires=request.form['evalPsicologResult']
        evpsicoreq=request.form['evalPsicometReq']
        evpsicop=request.form['evalPsicometPresene']
        evpsicores=request.form['evalPsicometResult']
        evtecreq=request.form['evalTecnicaReq']
        evtecp=request.form['evalTecnicaPresen']
        evtecres=request.form['evalTecnicaResul']
        evcreq=request.form['evalConocReq']
        evcp=request.form['evalConocPresen']
        evcres=request.form['evalConocResult']
        entfreq=request.form['entrevFinalReq']
        entfp=request.form['entrevFinalPresen']
        entfres=request.form['entrevFinalResul']
        cursor.execute('update candidato set CURP=%s, RFC=%s, nombre=%s, domCalle=%s, domNumExtInt=%s, domColonia=%s, tel1=%s, tel2=%s, correoE=%s, edad=%s, sexo=%s, idEstadoCivil=%s, idEscolaridad=%s, idGradoAvance=%s, idCarrera=%s, entrevSelecReq=%s, entrevSelecPresen=%s, entrevSelecresult=%s, evalMedicaReq=%s, evalMedicaPresen=%s, evalMedicaResult=%s, evalPsicolgReq=%s, evalPsicologPresen=%s, evalPsicologResult=%s, evalPsicometReq=%s, evalPsicometPresene=%s, evalPsicometResult=%s, evalTecnicaReq=%s, evalTecnicaPresen=%s, evalPsicometResult=%s, evalConocReq=%s, evalConocPresen=%s, evalConocResult=%s, entrevFinalReq=%s, entrevFinalPresen=%s, entrevFinalResul=%s where idCandidato= %s', (curp, rfc, nom, calle, numex, col, tel1, tel2, correo, edad, sexo, idec, ide, idga, idc, entsreq, entsp, entsres, evmedreq, evmedp, evmedres, evpsicoreq, evpsip, evpsires, evpsicoreq, evpsicop, evpsicores, evtecreq, evtecp, evtecres, evcreq, evcp, evcres, entfreq,entfp, entfres, id))
        conn.commit()
    return redirect(url_for('candidatos'))

@app.route('/candidatos_borrar/<string:id>')
def candidatos_borrar(id):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute('delete from candidato where idCandidato = {0}'.format(id))
    conn.commit()
    return redirect(url_for('candidatos'))

@app.route('/aprobar/<string:id>')
def aprobar_candidato(id):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute('select idVacante from candidato where idCandidato=%s', (id))
    vacante=cursor.fetchone()
    cursor.execute('update vacante set candidatoSelecc=%s where idVacante=%s', (id,vacante[0]))
    cursor.execute('update candidato set estado = "aceptado" WHERE idCandidato = %s', (id,))
    conn.commit()
    return redirect(url_for('vacante'))

@app.route('/rechazar/<string:id>')
def rechazar_candidato(id):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute('select nombre from candidato where idCandidato=%s', (id))
    nombre=cursor.fetchone()
    cursor.execute('insert into candidatosR(nombre) values(%s)', (nombre))
    cursor.execute('delete from candidato where idCandidato = {0}'.format(id))
    conn.commit()
    return redirect(url_for('candidatos'))

#Registro candidato
@app.route('/registrar/<string:id>')
def registrar(id):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute('select idVacante from vacante where idVacante = %s', (id))
    dato  = cursor.fetchall()
    cursor.execute('SELECT puesto.nomPuesto FROM puesto JOIN vacante ON puesto.idPuesto = vacante.idPuesto WHERE vacante.idVacante = %s', (id))
    datou = cursor.fetchone()

    cursor.execute('SELECT DISTINCT descripcion FROM sexo')  
    opciones_publicada = [row[0] for row in cursor.fetchall()] 
    cursor.execute('SELECT DISTINCT descripcion FROM estado_civil')  
    opciones_publicada2 = [row[0] for row in cursor.fetchall()]  
    cursor.execute('SELECT DISTINCT descripcion FROM escolaridad')  
    opciones_publicada3 = [row[0] for row in cursor.fetchall()]  
    cursor.execute('SELECT DISTINCT descripcion FROM grado_avance')  
    opciones_publicada4 = [row[0] for row in cursor.fetchall()]  
    cursor.execute('SELECT DISTINCT descripcion FROM carrera')  
    opciones_publicada5 = [row[0] for row in cursor.fetchall()]  
    conn.close()
    return render_template("registrar.html", comentar=dato[0],datou=datou,opciones_publicada=opciones_publicada,opciones_publicada2=opciones_publicada2,opciones_publicada3=opciones_publicada3,opciones_publicada4=opciones_publicada4,opciones_publicada5=opciones_publicada5)

@app.route('/registrar_fagrega', methods=['POST'])
def registrar_fagrega():
    conn =conectar()
    cursor = conn.cursor()
    idr=0
    idp=0
    ide=0
    ides=0
    idg=0
    idc=0
    if request.method == 'POST':

        idvac =request.form['Vacante']
        cursor.execute('select idRequisicion from vacante where idVacante=%s', (idvac))
        idr=cursor.fetchone()
        cursor.execute('select idPuesto from vacante where idVacante=%s', (idvac))
        idp=cursor.fetchone()
        curp=request.form['curp']
        rfc=request.form['rfc']
        nombre=request.form['Nombre']
        calle=request.form['Calle']
        numero=request.form['Numero']
        colonia=request.form['Colonia']
        tel1=request.form['tel1']
        tel2=request.form['tel2']
        correo=request.form['Correo']
        edad=request.form['edad']
        sexo=request.form['sexo']
        estado=request.form['Estado']
        cursor.execute('select idEstadoCivil from estado_civil where descripcion=%s', (estado))
        ide=cursor.fetchone()
        escolaridad=request.form['Escolaridad']
        cursor.execute('select idEscolaridad from escolaridad where descripcion=%s', (escolaridad))
        ides=cursor.fetchone()
        grado=request.form['Grado']
        cursor.execute('select idGradoAvance from grado_avance where descripcion=%s', (grado))
        idg=cursor.fetchone()
        carrera=request.form['Carrera']
        cursor.execute('select idCarrera from carrera where descripcion=%s', (carrera))
        idc=cursor.fetchone()


        cursor.execute('insert into candidato (idVacante, idRequisicion, idPuesto, CURP, RFC, nombre, domCalle, domNumExtInt, domColonia, tel1, tel2, correoE, edad, sexo, idEstadoCivil, idEscolaridad, idGradoAvance, idCarrera) values (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)',
               (idvac, idr, idp, curp, rfc, nombre, calle, numero, colonia, tel1, tel2, correo, edad, sexo, ide, ides, idg, idc))
        conn.commit()
    return redirect(url_for('home'))

#Equipo 4 Sofia#
@app.route('/generar_documento_word/<int:idCandidato>', methods=['POST'])
def generar_documento_word(idCandidato):
        conn = conectar()
        cursor = conn.cursor()

        cursor.execute("""
        SELECT cand.idCandidato, req.nomSolicita AS Nombre_Solicitante, cand.nombre AS Nombre_Candidato, cand.rfc AS RFC_Candidato, cand.CURP, ec.descripcion AS Estado_Civil, cand.edad AS Edad, cand.sexo AS Sexo, cand.domCalle AS Calle, cand.domNumExtInt AS Numero_Ext, cand.domColonia AS Colonia, cand.idPuesto AS ID_Puesto, p.nomPuesto AS Nombre_Puesto, p.descripcionGeneral AS Descripcion_Puesto, p.jornada AS Jornada, p.remunMensual AS Remuneracion_Mensual, p.reqFisicos AS Requisitos_Fisicos, p.reqPsicologicos AS Requisitos_Psicologicos, req.tipoVacante AS Tipo_Vacante, cand.entrevFinalResul AS Resultado_Entrevista_Final 
        FROM requisicion req 
        JOIN candidato cand ON req.idRequisicion = cand.idRequisicion 
        JOIN puesto p ON cand.idPuesto = p.idPuesto 
        JOIN estado_civil ec ON cand.idEstadoCivil = ec.idEstadoCivil
        WHERE cand.idCandidato = %s
        """, (idCandidato,))

        resultados = cursor.fetchall()

        if not resultados:
            return "Datos erróneos o Incompletos"



        column_names = [column[0] for column in cursor.description]
        valores_por_columna = [[row[i] for row in resultados] for i in range(len(cursor.description))]
        idCandidato = valores_por_columna[column_names.index("idCandidato")]
        Nombre_Solicitante = valores_por_columna[column_names.index("Nombre_Solicitante")]
        Nombre_Candidato = valores_por_columna[column_names.index("Nombre_Candidato")]
        RFC_Candidato = valores_por_columna[column_names.index("RFC_Candidato")]
        CURP = valores_por_columna[column_names.index("CURP")]
        Estado_Civil = valores_por_columna[column_names.index("Estado_Civil")]
        Edad = valores_por_columna[column_names.index("Edad")]
        Sexo = valores_por_columna[column_names.index("Sexo")]
        Calle = valores_por_columna[column_names.index("Calle")]
        Numero_Ext = valores_por_columna[column_names.index("Numero_Ext")]
        Colonia = valores_por_columna[column_names.index("Colonia")]
        ID_Puesto = valores_por_columna[column_names.index("ID_Puesto")]
        Nombre_Puesto = valores_por_columna[column_names.index("Nombre_Puesto")]
        Descripcion_Puesto = valores_por_columna[column_names.index("Descripcion_Puesto")]
        Jornada = valores_por_columna[column_names.index("Jornada")]
        Remuneracion_Mensual = valores_por_columna[column_names.index("Remuneracion_Mensual")]
        Requisitos_Fisicos = valores_por_columna[column_names.index("Requisitos_Fisicos")]
        Requisitos_Psicologicos = valores_por_columna[column_names.index("Requisitos_Psicologicos")]
        Tipo_Vacante = valores_por_columna[column_names.index("Tipo_Vacante")]
        Resultado_Entrevista_Final = valores_por_columna[column_names.index("Resultado_Entrevista_Final")]

        dir_actual = os.path.dirname(__file__)
        plantilla = os.path.join(dir_actual, 'static', 'plantilla.docx')
        document = Document(plantilla)
        for resultado in resultados:
         for paragraph in document.paragraphs:
            for field_name, field_value in zip(column_names, resultado):
                marcador = f"[{field_name}]"
                if marcador in paragraph.text:
                    paragraph.text = paragraph.text.replace(marcador, str(field_value))
                if "[Nombre_Solicitante]" in paragraph.text:
                    paragraph.text = paragraph.text.replace("[Nombre_Solicitante]", str(Nombre_Solicitante[0]))
                if "[Nombre_Candidato]" in paragraph.text:
                    paragraph.text = paragraph.text.replace("[Nombre_Candidato]", str(Nombre_Candidato[0]))
                if "[RFC_Candidato]" in paragraph.text:
                 paragraph.text = paragraph.text.replace("[RFC_Candidato]", str(RFC_Candidato[0]))
                if "[CURP]" in paragraph.text:
                 paragraph.text = paragraph.text.replace("[CURP]", str(CURP[0]))
                if "[Estado_Civil]" in paragraph.text:
                 paragraph.text = paragraph.text.replace("[Estado_Civil]", str(Estado_Civil[0]))
                if "[Edad]" in paragraph.text:
                    paragraph.text = paragraph.text.replace("[Edad]", str(Edad[0]))
                if "[Sexo]" in paragraph.text:
                    paragraph.text = paragraph.text.replace("[Sexo]", str(Sexo[0]))
                if "[Calle]" in paragraph.text:
                 paragraph.text = paragraph.text.replace("[Calle]", str(Calle[0]))
                if "[Numero_Ext]" in paragraph.text:
                 paragraph.text = paragraph.text.replace("[Numero_Ext]", str(Numero_Ext[0]))                 
                if "[Colonia]" in paragraph.text:
                    paragraph.text = paragraph.text.replace("[Colonia]", str(Colonia[0]))
                if "[ID_Puesto]" in paragraph.text:
                 paragraph.text = paragraph.text.replace("[ID_Puesto]", str(ID_Puesto[0]))
                if "[Nombre_Puesto]" in paragraph.text:
                 paragraph.text = paragraph.text.replace("[Nombre_Puesto]", str(Nombre_Puesto[0]))
                if "[Descripcion_Puesto]" in paragraph.text:
                    paragraph.text = paragraph.text.replace("[Descripcion_Puesto]", str(Descripcion_Puesto[0]))
                if "[Jornada]" in paragraph.text:
                 paragraph.text = paragraph.text.replace("[Jornada]", str(Jornada[0]))
                if "[Remuneracion_Mensual]" in paragraph.text:
                 paragraph.text = paragraph.text.replace("[Remuneracion_Mensual]", str(Remuneracion_Mensual[0]))
                if "[Requisitos_Fisicos]" in paragraph.text:
                    paragraph.text = paragraph.text.replace("[Requisitos_Fisicos]", str(Requisitos_Fisicos[0]))
                if "[Requisitos_Psicologicos]" in paragraph.text:
                 paragraph.text = paragraph.text.replace("[Requisitos_Psicologicos]", str(Requisitos_Psicologicos[0]))
                if "[Tipo_Vacante ]" in paragraph.text:
                 paragraph.text = paragraph.text.replace("[Tipo_Vacante]", str(Tipo_Vacante[0]))
                if "[Resultado_Entrevista_Final]" in paragraph.text:
                    paragraph.text = paragraph.text.replace("[Resultado_Entrevista_Final]", str(Resultado_Entrevista_Final[0]))
        cursor.close()
        conn.close()
        pythoncom.CoInitialize()
        document.save('contrato.docx')
        docx2pdf.convert("contrato.docx")
        return send_file('contrato.pdf', as_attachment=True)

@app.route('/candidato')
def candidato():
    conn =conectar()
    cursor = conn.cursor()
    cursor.execute('SELECT c.idCandidato, c.nombre, p.nomPuesto ,c.estado FROM candidato c INNER JOIN puesto p ON c.idPuesto = p.idPuesto WHERE c.estado = "aceptado"')
    datos = cursor.fetchall()
    return render_template("candidato.html", comentarios = datos)

#Equipo 2 Gael#  


@app.route('/form')
def form():
    return render_template("form.html")

@app.route('/crear_anuncio', methods=['POST'])
def crear_anuncio():
    if request.method == 'POST':
        # Obtener los datos del formulario
        id_vacante = request.form['id_vacante']

        folio = request.form['folio']
        puesto = request.form['puesto']
        personal = request.form['personal']
        medios = ', '.join(request.form.getlist('medios'))
        fecha_inicio = request.form['fechaInicio']
        fecha_finalizacion = request.form['fechaFinalizacion']
        fecha_elegida = request.form['fechaElegida'] if fecha_finalizacion != 'N/A' else 'N/A'

        # Crear un nuevo documento de Word
        doc = Document()
        doc.add_heading('Anuncio de Vacante', level=1)

        # Agregar los datos del formulario al documento
        doc.add_paragraph(f'ID Vacante: {id_vacante}')
        doc.add_paragraph(f'Folio: {folio}')
        doc.add_paragraph(f'Puesto: {puesto}')
        doc.add_paragraph(f'Personal: {personal}')
        doc.add_paragraph(f'Medios: {medios}')
        doc.add_paragraph(f'Fecha de inicio: {fecha_inicio}')
        doc.add_paragraph(f'Fecha de finalización: {fecha_finalizacion}')

        # Guardar el documento en un archivo
        doc.save('anuncio.docx')

        # Retornar una respuesta indicando que el archivo se ha creado exitosamente
        return render_template ("layout2.html") 
    

    
@app.route('/trabajadores')
def trabajadores():
    conn =conectar()
    cursor = conn.cursor()
    cursor.execute('select idTrabajador, NombreTrab from trabajadores order by idTrabajador')
    datos = cursor.fetchall()
    return render_template("trabajadores.html", trab = datos)




@app.route('/trabajador_agregar')
def trabajador_agregar():
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute('select idArea, descripcion from area ')
    datos1 = cursor.fetchall()

    

    cursor.execute('select idEstadoCivil, descripcion from estado_civil ')
    datos2 = cursor.fetchall()

    cursor.execute('select idEscolaridad, descripcion from escolaridad ')
    datos3 = cursor.fetchall()

    cursor.execute('select idGradoAvance, descripcion from grado_avance ')
    datos4 = cursor.fetchall()

    cursor.execute('select idCarrera, descripcion from carrera ')
    datos5 = cursor.fetchall()

    cursor.execute('select idIdioma, descripcion from idioma ')
    datos6 = cursor.fetchall()

    cursor.execute('select idHabilidad, descripcion from habilidad ')
    datos7 = cursor.fetchall()

    cursor.execute('select idPuesto, nomPuesto from puesto ')
    datos8= cursor.fetchall()

    return render_template("trabajador_agr.html", catArea=datos1, catEdoCivil=datos2, catEscolaridad=datos3,
                           catGradoAvance=datos4, catCarrera=datos5, catIdioma=datos6, catHabilidad=datos7, catPuesto=datos8)

@app.route('/trabajador_fagrega', methods=['POST'])
def trabajador_fagrega():
    if request.method == 'POST':

        if  'idArea' in request.form:
            idAr = request.form['idArea']
        else:
            idAr = '1'
        if 'idEstadoCivil' in request.form:
            idEC = request.form['idEstadoCivil']
        else:
            idEC = '1'
        if 'idPuesto' in request.form:
            idPu = request.form['idPuesto']
        else:
            idPu = '1'
        if 'idEscolaridad' in request.form:
            idEs = request.form['idEscolaridad']
        else:
            idEs = '1'
        if 'idGradoAvance' in request.form:
            idGA = request.form['idGradoAvance']
        else:
            idGA = '1'
        if 'idCarrera' in request.form:
            idCa = request.form['idCarrera']
        else:
            idCa = '1'
        if 'sexo' in request.form:
            sex = request.form['sexo']
        else:
            sex = '1'
        
        nomT = request.form['nomtrab']
        eda = request.form['edad']



    conn = conectar()
    cursor = conn.cursor()
    cursor.execute('insert into trabajadores (idArea, idPuesto, nombreTrab, edad,sexo,idEstadoCivil,idEscolaridad,idGradoAvance,idCarrera'') values (%s,%s,%s,%s,%s,%s,%s,%s,%s)',( idAr, idPu, nomT, eda, sex ,idEC, idEs, idGA, idCa))
    conn.commit()

    cursor.execute('select idTrabajador from trabajadores where idTrabajador=(select max(idTrabajador) from trabajadores) ')
    dato = cursor.fetchall()
    idpue = dato[0]
    idP = idpue[0]


    return redirect(url_for('trabajadores'))

@app.route('/trabajadores_borrar/<string:idP>')
def trabajadores_borrar(idP):
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute('delete from trabajadores where idTrabajador = %s',(idP))
    conn.commit()
  
    return redirect(url_for('trabajadores'))


@app.route('/trabajador_editar/<string:idP>')
def trabajador_editar(idP):
    conn = conectar()
    cursor = conn.cursor()

    cursor.execute('select idTrabajador, idArea, idPuesto, nombreTrab, edad, sexo, idEstadoCivil, idEscolaridad, idGradoAvance, idCarrera from trabajadores where idTrabajador = %s', (idP))
    dato = cursor.fetchall()

    cursor.execute('select idArea, descripcion from area')
    datos1 = cursor.fetchall()

    cursor.execute('select idPuesto, nomPuesto from puesto')
    datos2 = cursor.fetchall()

    cursor.execute('select idEstadoCivil, descripcion from estado_civil')
    datos3 = cursor.fetchall()

    cursor.execute('select idEscolaridad, descripcion from escolaridad')
    datos4 = cursor.fetchall()

    cursor.execute('select idGradoAvance, descripcion from grado_avance')
    datos5 = cursor.fetchall()

    cursor.execute('select idCarrera, descripcion from carrera')
    datos6 = cursor.fetchall()



    cursor.execute('select a.idArea, a.descripcion from area a, trabajadores b where a.idArea = b.idArea and b.idTrabajador = %s', (idP))
    datos9 = cursor.fetchall()

    cursor.execute('select a.idPuesto, a.nomPuesto from puesto a, trabajadores b where a.idPuesto = b.idPuesto and b.idTrabajador = %s', (idP))
    datos10 = cursor.fetchall()

    cursor.execute('select a.idEstadoCivil, a.descripcion from estado_civil a, trabajadores b where a.idEstadoCivil = b.idEstadoCivil and b.idTrabajador = %s', (idP))
    datos11 = cursor.fetchall()

    cursor.execute('select a.idEscolaridad, a.descripcion from escolaridad a, trabajadores b where a.idEscolaridad = b.idEscolaridad and b.idTrabajador = %s', (idP))
    datos12 = cursor.fetchall()

    cursor.execute('select a.idGradoAvance, a.descripcion from grado_avance a, trabajadores b where a.idGradoAvance = b.idGradoAvance and b.idTrabajador = %s', (idP))
    datos13 = cursor.fetchall()

    cursor.execute('select a.idCarrera, a.descripcion from carrera a, trabajadores b where a.idCarrera = b.idCarrera and b.idTrabajador = %s', (idP))
    datos14 = cursor.fetchall()


    return render_template("trabajador_edi.html", dat=dato[0], catArea=datos1, catPuesto=datos2, catEdoCivil=datos3, catEscolaridad=datos4,
                           catGradoAvance=datos5, catCarrera=datos6, Area=datos9[0], Puesto=datos10[0], EdoCivil=datos11[0], Escolaridad=datos12[0], GradoAvance=datos13[0],
                           Carrera=datos14[0])


@app.route('/trabajador_fedita/<string:idP>', methods=['POST'])
def trabajador_fedita(idP):
    if request.method == 'POST':
        idAr = request.form['idArea']
        idPu = request.form['idPuesto']
        nomT = request.form['nomtrab']
        eda = request.form['edad']
        sex = request.form['sexo']
        idEC = request.form['idEstadoCivil']
        idEs = request.form['idEscolaridad']
        idGA = request.form['idGradoAvance']
        idCa = request.form['idCarrera']
        
        conn = conectar()
        cursor = conn.cursor()
        
        cursor.execute(
            'UPDATE trabajadores SET idArea = %s, idPuesto = %s, NombreTrab = %s, edad = %s, sexo = %s, '
            'idEstadoCivil = %s, idEscolaridad = %s, idGradoAvance = %s, idCarrera = %s WHERE idTrabajador = %s',
            (idAr, idPu, nomT, eda, sex, idEC, idEs, idGA, idCa, idP)
        )
        conn.commit()
        
        
        return redirect(url_for('trabajadores'))
    
    
@app.route('/trabajo_fdetalle/<string:idP>', methods=['GET'])
def trabajador_fdetalle(idP):
    conn = conectar()
    cursor = conn.cursor()

    cursor.execute('select idTrabajador, nombreTrab from trabajadores order by idTrabajador')
    datos = cursor.fetchall()

    cursor.execute('SELECT idTrabajador, idArea, idPuesto, NombreTrab, '
               'edad, sexo, idEstadoCivil, idEscolaridad, idGradoAvance, idCarrera '
               'FROM trabajadores WHERE idTrabajador = %s', (idP,))
    dato = cursor.fetchall()


    cursor.execute('select a.idArea, a.descripcion from area a, trabajadores b where a.idArea = b.idArea and b.idTrabajador = %s', (idP))
    datos1 = cursor.fetchall()

    cursor.execute('select a.idPuesto, a.nomPuesto from puesto a, trabajadores b where a.idPuesto = b.idPuesto and b.idTrabajador = %s', (idP))
    datos2 = cursor.fetchall()

    cursor.execute('select a.idEstadoCivil, a.descripcion from estado_civil a, trabajadores b where a.idEstadoCivil = b.idEstadoCivil and b.idTrabajador = %s', (idP))
    datos3 = cursor.fetchall()

    cursor.execute('select a.idEscolaridad, a.descripcion from escolaridad a, trabajadores b where a.idEscolaridad = b.idEscolaridad and b.idTrabajador = %s', (idP))
    datos4 = cursor.fetchall()

    cursor.execute('select a.idGradoAvance, a.descripcion from grado_avance a, trabajadores b where a.idGradoAvance = b.idGradoAvance and b.idTrabajador = %s', (idP))
    datos5 = cursor.fetchall()

    cursor.execute('select a.idCarrera, a.descripcion from carrera a, trabajadores b where a.idCarrera = b.idCarrera and b.idTrabajador = %s', (idP))
    datos6 = cursor.fetchall()

   
    return render_template("trab_detalle.html", pue = datos, dat=dato[0], catArea=datos1[0],catPuesto=datos2[0], catEdoCivil=datos3[0], catEscolaridad=datos4[0],
                           catGradoAvance=datos5[0], catCarrera=datos6[0])

@app.route('/curso_agrOp2')
def curso_agrOp2():
    conn = conectar()
    cursor = conn.cursor()
    cursor.execute('select id_curso, nombre from cursos ')
    datos1 = cursor.fetchall()

    cursor.execute('select idTrabajador, NombreTrab from trabajadores ')
    datos2 = cursor.fetchall()


    return render_template("addCurso.html", catCurso=datos1, catTrab=datos2)

@app.route('/addCurso', methods=['POST'])
def curso_fagrega():
    if request.method == 'POST':
        idT = request.form.get('idTrab')
        idCu = request.form.get('idCurso')

        try:
            conn = conectar()
            cursor = conn.cursor()
            cursor.execute('insert into curso_has_empleados (id_curso,id_empleado) values (%s,%s)', (idT, idCu))
            conn.commit()
            cursor.close()
            conn.close()
        except pymysql.MySQLError as e:
            print(f"Error: {e}")
            # Aquí podrías agregar un manejo de errores más sofisticado, como registrar el error o notificar al usuario

        return redirect(url_for('trabajadores'))

def get_db_connection():
    return conectar()

@app.route('/CursosTrab/<int:idTrabajador>', methods=['GET'])
def CursosTrab(idTrabajador):
    conn = get_db_connection()
    cursor = conn.cursor()

    # Obtener información del trabajador
    cursor.execute('SELECT NombreTrab FROM trabajadores WHERE idTrabajador = %s', (idTrabajador,))
    trabajador = cursor.fetchone()

    # Obtener los cursos asociados al trabajador
    cursor.execute('''
    SELECT c.id_curso, c.nombre, c.duracion, b.edicion
    FROM cursos c
    JOIN curso_has_empleados tc ON c.id_curso = tc.id_curso
    JOIN curso_has_aparicion b ON c.id_curso = b.id_curso
    WHERE tc.id_empleado = %s
    ''', (idTrabajador,))
    cursos = cursor.fetchall()

  
    cursor.close()
    conn.close()

    return render_template("capacitacionesT.html", trabajador=trabajador, cursos=cursos)

@app.route('/generar_documento_word/<int:idCandidato>', methods=['POST'])
def generar_documento_word(idCandidato):
        conn = pymysql.connect(host='localhost', user='root', passwd='', db='rh3')
        cursor = conn.cursor()

        cursor.execute("""
        SELECT cand.idCandidato, req.nomSolicita AS Nombre_Solicitante, cand.nombre AS Nombre_Candidato, cand.rfc AS RFC_Candidato, cand.CURP, ec.descripcion AS Estado_Civil, cand.edad AS Edad, cand.sexo AS Sexo, cand.domCalle AS Calle, cand.domNumExtInt AS Numero_Ext, cand.domColonia AS Colonia, cand.idPuesto AS ID_Puesto, p.nomPuesto AS Nombre_Puesto, p.descripcionGeneral AS Descripcion_Puesto, p.jornada AS Jornada, p.remunMensual AS Remuneracion_Mensual, p.reqFisicos AS Requisitos_Fisicos, p.reqPsicologicos AS Requisitos_Psicologicos, req.tipoVacante AS Tipo_Vacante, cand.entrevFinalResul AS Resultado_Entrevista_Final 
        FROM requisicion req 
        JOIN candidato cand ON req.idRequisicion = cand.idRequisicion 
        JOIN puesto p ON cand.idPuesto = p.idPuesto 
        JOIN estado_civil ec ON cand.idEstadoCivil = ec.idEstadoCivil
        WHERE cand.idCandidato = %s
        """, (idCandidato,))

        resultados = cursor.fetchall()

        if not resultados:
            return "Datos erróneos o Incompletos"



        column_names = [column[0] for column in cursor.description]
        valores_por_columna = [[row[i] for row in resultados] for i in range(len(cursor.description))]
        idCandidato = valores_por_columna[column_names.index("idCandidato")]
        Nombre_Solicitante = valores_por_columna[column_names.index("Nombre_Solicitante")]
        Nombre_Candidato = valores_por_columna[column_names.index("Nombre_Candidato")]
        RFC_Candidato = valores_por_columna[column_names.index("RFC_Candidato")]
        CURP = valores_por_columna[column_names.index("CURP")]
        Estado_Civil = valores_por_columna[column_names.index("Estado_Civil")]
        Edad = valores_por_columna[column_names.index("Edad")]
        Sexo = valores_por_columna[column_names.index("Sexo")]
        Calle = valores_por_columna[column_names.index("Calle")]
        Numero_Ext = valores_por_columna[column_names.index("Numero_Ext")]
        Colonia = valores_por_columna[column_names.index("Colonia")]
        ID_Puesto = valores_por_columna[column_names.index("ID_Puesto")]
        Nombre_Puesto = valores_por_columna[column_names.index("Nombre_Puesto")]
        Descripcion_Puesto = valores_por_columna[column_names.index("Descripcion_Puesto")]
        Jornada = valores_por_columna[column_names.index("Jornada")]
        Remuneracion_Mensual = valores_por_columna[column_names.index("Remuneracion_Mensual")]
        Requisitos_Fisicos = valores_por_columna[column_names.index("Requisitos_Fisicos")]
        Requisitos_Psicologicos = valores_por_columna[column_names.index("Requisitos_Psicologicos")]
        Tipo_Vacante = valores_por_columna[column_names.index("Tipo_Vacante")]
        Resultado_Entrevista_Final = valores_por_columna[column_names.index("Resultado_Entrevista_Final")]

        dir_actual = os.path.dirname(__file__)
        plantilla = os.path.join(dir_actual, 'static', 'plantilla.docx')
        document = Document(plantilla)
        for resultado in resultados:
         for paragraph in document.paragraphs:
            for field_name, field_value in zip(column_names, resultado):
                marcador = f"[{field_name}]"
                if marcador in paragraph.text:
                    paragraph.text = paragraph.text.replace(marcador, str(field_value))
                if "[Nombre_Solicitante]" in paragraph.text:
                    paragraph.text = paragraph.text.replace("[Nombre_Solicitante]", str(Nombre_Solicitante[0]))
                if "[Nombre_Candidato]" in paragraph.text:
                    paragraph.text = paragraph.text.replace("[Nombre_Candidato]", str(Nombre_Candidato[0]))
                if "[RFC_Candidato]" in paragraph.text:
                 paragraph.text = paragraph.text.replace("[RFC_Candidato]", str(RFC_Candidato[0]))
                if "[CURP]" in paragraph.text:
                 paragraph.text = paragraph.text.replace("[CURP]", str(CURP[0]))
                if "[Estado_Civil]" in paragraph.text:
                 paragraph.text = paragraph.text.replace("[Estado_Civil]", str(Estado_Civil[0]))
                if "[Edad]" in paragraph.text:
                    paragraph.text = paragraph.text.replace("[Edad]", str(Edad[0]))
                if "[Sexo]" in paragraph.text:
                    paragraph.text = paragraph.text.replace("[Sexo]", str(Sexo[0]))
                if "[Calle]" in paragraph.text:
                 paragraph.text = paragraph.text.replace("[Calle]", str(Calle[0]))
                if "[Numero_Ext]" in paragraph.text:
                 paragraph.text = paragraph.text.replace("[Numero_Ext]", str(Numero_Ext[0]))                 
                if "[Colonia]" in paragraph.text:
                    paragraph.text = paragraph.text.replace("[Colonia]", str(Colonia[0]))
                if "[ID_Puesto]" in paragraph.text:
                 paragraph.text = paragraph.text.replace("[ID_Puesto]", str(ID_Puesto[0]))
                if "[Nombre_Puesto]" in paragraph.text:
                 paragraph.text = paragraph.text.replace("[Nombre_Puesto]", str(Nombre_Puesto[0]))
                if "[Descripcion_Puesto]" in paragraph.text:
                    paragraph.text = paragraph.text.replace("[Descripcion_Puesto]", str(Descripcion_Puesto[0]))
                if "[Jornada]" in paragraph.text:
                 paragraph.text = paragraph.text.replace("[Jornada]", str(Jornada[0]))
                if "[Remuneracion_Mensual]" in paragraph.text:
                 paragraph.text = paragraph.text.replace("[Remuneracion_Mensual]", str(Remuneracion_Mensual[0]))
                if "[Requisitos_Fisicos]" in paragraph.text:
                    paragraph.text = paragraph.text.replace("[Requisitos_Fisicos]", str(Requisitos_Fisicos[0]))
                if "[Requisitos_Psicologicos]" in paragraph.text:
                 paragraph.text = paragraph.text.replace("[Requisitos_Psicologicos]", str(Requisitos_Psicologicos[0]))
                if "[Tipo_Vacante ]" in paragraph.text:
                 paragraph.text = paragraph.text.replace("[Tipo_Vacante]", str(Tipo_Vacante[0]))
                if "[Resultado_Entrevista_Final]" in paragraph.text:
                    paragraph.text = paragraph.text.replace("[Resultado_Entrevista_Final]", str(Resultado_Entrevista_Final[0]))
        cursor.close()
        conn.close()
        pythoncom.CoInitialize()
        document.save('contrato.docx')
        docx2pdf.convert("contrato.docx")
        return send_file('contrato.pdf', as_attachment=True)






if __name__ == "__main__":
    app.run(debug=True)
